import os
import markdown
from bs4 import BeautifulSoup
from docx import Document
from typing import Union
from llama_index.core import (
    VectorStoreIndex,
    StorageContext,
)
import lancedb
from llama_parse import LlamaParse
from llama_parse.utils import ResultType
from llama_index.vector_stores.lancedb import LanceDBVectorStore
from llama_index.core.indices.base import BaseIndex
from llama_index.llms.openai import OpenAI

# Environment Variables
VECTOR_STORE_URI = os.getenv("VS_URI", "./lancedb")


def drop_project_table(project_name: str):
    """
    Delete the index from the data room path
    """
    db = lancedb.connect(uri=VECTOR_STORE_URI)
    db.drop_table(project_name)


def get_index_from_project(project_name: str) -> Union[VectorStoreIndex, BaseIndex]:
    """
    Get the index from the data room path
    """
    vector_store = LanceDBVectorStore(
        uri=VECTOR_STORE_URI,
        mode="read",
        table_name=project_name,
    )
    index = VectorStoreIndex.from_vector_store(vector_store)
    return index


# Document Ingestion and Vector Index Creation
async def ingest_documents(data_room_path: str) -> Union[VectorStoreIndex, BaseIndex]:
    """
    Ingest documents from a dataroom path and create a vector index
    """
    table_name = os.path.basename(os.path.normpath(data_room_path))

    parser = LlamaParse(result_type=ResultType.MD)  # return markdown

    file_paths = [
        os.path.join(data_room_path, path)
        for path in os.listdir(data_room_path)
        if os.path.isfile(os.path.join(data_room_path, path))
    ]

    documents = await parser.aload_data(file_paths)  # type: ignore

    llm = OpenAI(temperature=0.1, model="gpt-4o")

    ## TODO in the future we will want to look at text hash to determine if the document has changed
    ## and only update the changed documents, not the entire vector index
    vector_store = LanceDBVectorStore(
        uri=VECTOR_STORE_URI,
        mode="overwrite",
        table_name=table_name,
    )
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    index = VectorStoreIndex.from_documents(
        documents, storage_context=storage_context, llm=llm
    )

    return index


def markdown_to_word(markdown_path: str):
    """
    Convert a markdown file to a word document
    """
    # Generate the Word document path
    base_path = os.path.splitext(markdown_path)[0]
    docx_path = f"{base_path}.docx"

    # Read the markdown file
    with open(markdown_path, "r", encoding="utf-8") as md_file:
        md_content = md_file.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)

    # Parse HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # Create a new Word document
    doc = Document()

    # Convert HTML elements to Word elements
    for element in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol"]):
        if element.name == "p":
            doc.add_paragraph(element.text)
        elif element.name.startswith("h"):
            level = int(element.name[1])
            doc.add_heading(element.text, level=level)
        elif element.name in ["ul", "ol"]:
            for li in element.find_all("li"):
                doc.add_paragraph(
                    li.text,
                    style="List Bullet" if element.name == "ul" else "List Number",
                )

    # Save the document
    doc.save(docx_path)
    print(f"Converted {markdown_path} to {docx_path}")
